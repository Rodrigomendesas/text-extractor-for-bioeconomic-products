"""Text extraction utilities for various file formats."""

import logging
import io
import base64
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
import pdfplumber
import fitz  # PyMuPDF
from PIL import Image
import docx  # python-docx for DOCX files

from src.models.visual_element import VisualElement, VisualElementType

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extracts text and visual elements from various file formats."""

    def __init__(self):
        """Initialize the text extractor."""
        self.supported_extensions = {'.pdf', '.doc', '.docx'}
        self.extract_images = True  # Flag to control image extraction

    def extract_from_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text and visual elements from a file.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with extracted text, visual elements, and metadata
        """
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = file_path.suffix.lower()

        if extension == '.pdf':
            return self._extract_from_pdf(file_path)
        elif extension in ['.doc', '.docx']:
            return self._extract_from_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")

    def _extract_from_pdf(self, pdf_path: Path) -> Dict[str, Any]:
        """
        Extract text and visual elements from PDF using multiple methods for best results.

        Args:
            pdf_path: Path to the PDF file

        Returns:
            Dictionary with extracted text, visual elements, and metadata
        """
        logger.info(f"Extracting content from PDF: {pdf_path}")

        # Try multiple extraction methods for text
        methods = [
            ('pdfplumber', self._extract_with_pdfplumber),
            ('pymupdf', self._extract_with_pymupdf),
            ('pypdf2', self._extract_with_pypdf2)
        ]

        best_result = {'text': '', 'method': 'none', 'page_count': 0}

        for method_name, method_func in methods:
            try:
                result = method_func(pdf_path)
                if len(result['text']) > len(best_result['text']):
                    result['method'] = method_name
                    best_result = result
                    logger.info(f"Method {method_name} extracted {len(result['text'])} characters")
            except Exception as e:
                logger.warning(f"Method {method_name} failed: {e}")
                continue

        if not best_result['text']:
            raise ValueError("Could not extract text from PDF using any method")

        logger.info(f"Best extraction method: {best_result['method']}")

        # Extract visual elements if enabled
        visual_elements = []
        if self.extract_images:
            try:
                visual_elements = self._extract_visual_elements(pdf_path)
                logger.info(f"Extracted {len(visual_elements)} visual elements from PDF")
                best_result['visual_elements'] = visual_elements
            except Exception as e:
                logger.warning(f"Visual element extraction failed: {e}")
                best_result['visual_elements'] = []
        else:
            best_result['visual_elements'] = []

        return best_result

    def _extract_with_pdfplumber(self, pdf_path: Path) -> Dict[str, str]:
        """Extract text using pdfplumber (good for tables and complex layouts)."""
        text_parts = []
        page_count = 0

        with pdfplumber.open(pdf_path) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return {
            'text': '\n\n'.join(text_parts),
            'page_count': page_count
        }

    def _extract_with_pymupdf(self, pdf_path: Path) -> Dict[str, str]:
        """Extract text using PyMuPDF (fastest method)."""
        text_parts = []

        doc = fitz.open(pdf_path)
        page_count = len(doc)

        for page_num in range(page_count):
            page = doc.load_page(page_num)
            page_text = page.get_text()
            if page_text:
                text_parts.append(page_text)

        doc.close()

        return {
            'text': '\n\n'.join(text_parts),
            'page_count': page_count
        }

    def _extract_with_pypdf2(self, pdf_path: Path) -> Dict[str, str]:
        """Extract text using PyPDF2 (fallback method)."""
        text_parts = []

        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            page_count = len(pdf_reader.pages)

            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

        return {
            'text': '\n\n'.join(text_parts),
            'page_count': page_count
        }

    def _extract_visual_elements(self, pdf_path: Path) -> List[VisualElement]:
        """
        Extract images, charts, and tables from PDF using PyMuPDF.

        Args:
            pdf_path: Path to the PDF file

        Returns:
            List of VisualElement objects
        """
        visual_elements = []

        # Open the PDF with PyMuPDF
        doc = fitz.open(pdf_path)

        # Process each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)

            # Extract images
            image_list = page.get_images(full=True)

            for img_index, img_info in enumerate(image_list):
                try:
                    # Get the image
                    xref = img_info[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # Get image position on page
                    for img_rect in page.get_image_rects(xref):
                        position = (img_rect.x0, img_rect.y0, img_rect.x1, img_rect.y1)

                        # Create a PIL Image for metadata extraction
                        pil_img = Image.open(io.BytesIO(image_bytes))
                        width, height = pil_img.size

                        # Convert image to base64 for storage
                        base64_image = base64.b64encode(image_bytes).decode('utf-8')

                        # Create a VisualElement
                        element = VisualElement(
                            element_type=VisualElementType.IMAGE,
                            page_number=page_num + 1,
                            position=position,
                            content_base64=base64_image,
                            content_format=image_ext,
                            width=width,
                            height=height,
                            dpi=72  # Default DPI, could be refined
                        )

                        visual_elements.append(element)
                        break  # Use the first occurrence of this image on the page
                except Exception as e:
                    logger.warning(f"Failed to extract image {img_index} on page {page_num+1}: {e}")

            # Extract tables (basic detection based on rectangles and lines)
            # This is a simplified approach - for better table detection, consider using
            # specialized libraries like camelot-py or tabula-py
            try:
                # Look for potential tables based on rectangles and lines
                rects = page.search_for("", quads=True)  # Get all rectangles
                if len(rects) > 4:  # Simple heuristic: if we have several rectangles, it might be a table
                    # Get the page as an image
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    img_bytes = pix.tobytes("png")

                    # Convert to base64
                    base64_image = base64.b64encode(img_bytes).decode('utf-8')

                    # Create a VisualElement for the potential table
                    element = VisualElement(
                        element_type=VisualElementType.TABLE,
                        page_number=page_num + 1,
                        position=(0, 0, page.rect.width, page.rect.height),  # Full page for now
                        content_base64=base64_image,
                        content_format="png",
                        width=pix.width,
                        height=pix.height
                    )

                    visual_elements.append(element)
            except Exception as e:
                logger.warning(f"Failed to detect tables on page {page_num+1}: {e}")

        doc.close()
        return visual_elements

    def _extract_from_doc(self, doc_path: Path) -> Dict[str, Any]:
        """
        Extract text from DOC or DOCX files.

        Args:
            doc_path: Path to the DOC or DOCX file

        Returns:
            Dictionary with extracted text and metadata
        """
        logger.info(f"Extracting content from document: {doc_path}")

        extension = doc_path.suffix.lower()
        text = ""
        page_count = 0

        try:
            if extension == '.docx':
                # Use python-docx for DOCX files
                doc = docx.Document(doc_path)
                text_parts = []

                # Extract text from paragraphs
                for para in doc.paragraphs:
                    if para.text:
                        text_parts.append(para.text)

                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text:
                                row_text.append(cell.text)
                        if row_text:
                            text_parts.append(' | '.join(row_text))

                text = '\n'.join(text_parts)
                page_count = len(doc.sections)  # Approximate page count

                logger.info(f"Extracted {len(text)} characters from DOCX file")

            elif extension == '.doc':
                # For DOC files, we'll use a simple approach
                # Note: This is a simplified approach and may not work for all DOC files
                # A more robust solution would use libraries like pywin32 or antiword

                # Try to convert to text using external tools if available
                import subprocess
                import tempfile

                with tempfile.NamedTemporaryFile(suffix='.txt', delete=False) as temp_file:
                    temp_path = temp_file.name

                try:
                    # Try using textract if available
                    import textract
                    text = textract.process(str(doc_path)).decode('utf-8')
                    logger.info(f"Extracted {len(text)} characters from DOC file using textract")
                except (ImportError, Exception) as e:
                    logger.warning(f"Textract extraction failed: {e}")

                    # Fallback: Try using antiword if available
                    try:
                        result = subprocess.run(['antiword', str(doc_path)], 
                                               capture_output=True, text=True, check=True)
                        text = result.stdout
                        logger.info(f"Extracted {len(text)} characters from DOC file using antiword")
                    except (subprocess.SubprocessError, FileNotFoundError) as e:
                        logger.warning(f"Antiword extraction failed: {e}")

                        # Last resort: Try using catdoc if available
                        try:
                            result = subprocess.run(['catdoc', str(doc_path)], 
                                                   capture_output=True, text=True, check=True)
                            text = result.stdout
                            logger.info(f"Extracted {len(text)} characters from DOC file using catdoc")
                        except (subprocess.SubprocessError, FileNotFoundError) as e:
                            logger.warning(f"Catdoc extraction failed: {e}")
                            raise ValueError(f"Could not extract text from DOC file: {e}")

                # Estimate page count (very rough)
                page_count = max(1, len(text) // 3000)

            if not text:
                raise ValueError("No text extracted from document")

            return {
                'text': text,
                'page_count': page_count,
                'method': 'python-docx' if extension == '.docx' else 'external-tool',
                'visual_elements': []  # No visual elements for DOC/DOCX files for now
            }

        except Exception as e:
            logger.error(f"Failed to extract text from document {doc_path}: {e}")
            raise ValueError(f"Document extraction failed: {e}")

    def chunk_text(self, text: str, chunk_size: int = 3000, overlap: int = 200) -> List[str]:
        """
        Split text into chunks for processing.

        Args:
            text: Text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks

        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(min(100, chunk_size // 4)):
                    if end - i > start and text[end - i] in '.!?':
                        end = end - i + 1
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - overlap
            if start >= len(text):
                break

        return chunks
